from docx import Document

document = Document('demo.docx')


def getTexto():
    texto = []
    p1=[]
    p2 = []
    p3 = []
    p4 = []
    p5 = []
    for i in document.paragraphs:
        texto.append(i.text)
    lista = []
    for k in range(len(document.paragraphs)):
        try:
            if len(texto[k]) < 30:
                #print(texto[k])
                lista.append(texto[k] + "\n")
                if len(texto[k+1]) < 30:

                    if texto[k+1] not in lista:
                        lista.append(texto[k+1]+"\n")

            else:
                lista.append(texto[k])

            #print(texto[k])
                #lista.append(texto[k])

        except IndexError:
            pass


    #print(lista[2:11])
    lista.pop(1)
    lista.pop(2)
    for item in lista[1:11]:
        p1.append(item)
    #print(p1)
    for item in lista[11:20]:
        p2.append(item)

    for item in lista[20:34]:
        p3.append(item)

    for item in lista[38:43]:
        p4.append(item)

    for item in lista[43:83]:
        p5.append(item)


    td = []
    td.append(p1)
    td.append(p2)
    td.append(p3)
    td.append(p4)
    td.append(p5)
    print(td)
    for i in td:
        for j in i:
            print(j)

    '''for i in texto:
        if len(i)<20:
            x.append(i)
        else:
            x.insert(0,i)
    print(x)'''
def getStyleText():
    for p in document.paragraphs:
        print(p.style.name)
        #more common. Title,heading1,heading2,body text,normal, list paragraph	


def iter_heading(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading 1'): #or paragraph.style.name.startswith('Heading 2'):
            yield paragraph
def iter(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading 2'): #or paragraph.style.name.startswith('Heading 2'):
            yield paragraph




def main():
    for i in range(len(document.paragraphs)):
        print(i,"-",document.paragraphs[i].text,len(document.paragraphs[i].text))
        if len(document.paragraphs[i].text) < 30:
            print("\n")

    for i in range(10):
        print(document.paragraphs[i].text)

    print(document.paragraphs[4])
    print(document.paragraphs[4].runs[3].text) #pegar palavras com formatação diferente


    #getTexto()

main()